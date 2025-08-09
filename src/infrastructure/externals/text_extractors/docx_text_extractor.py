"""Word文書テキスト抽出実装。

Word文書（.docx）からテキストを抽出する。
"""

import io

from docx import Document as DocxDocument

from ....domain.externals import ExtractedText, TextExtractor


class DocxTextExtractor(TextExtractor):
    """Word文書テキスト抽出実装。

    python-docxライブラリを使用してWord文書からテキストを抽出する。
    """

    SUPPORTED_TYPES = {
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/msword",
    }

    async def extract_text(
        self, content: bytes, content_type: str
    ) -> ExtractedText:
        """Word文書からテキストを抽出する。

        Args:
            content: Word文書のバイナリデータ
            content_type: ファイルのMIMEタイプ

        Returns:
            抽出されたテキスト

        Raises:
            ValueError: サポートされていない形式の場合
            Exception: 文書処理でエラーが発生した場合
        """
        if not self.supports(content_type):
            raise ValueError(f"Unsupported content type: {content_type}")

        try:
            # バイトストリームから文書を読み込み
            doc_file = io.BytesIO(content)
            doc = DocxDocument(doc_file)

            # 段落からテキストを抽出
            text_parts = []
            paragraph_count = 0

            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
                    text_parts.append("\n")
                    paragraph_count += 1

            # テーブルからもテキストを抽出
            table_count = 0
            for table in doc.tables:
                table_count += 1
                for row in table.rows:
                    row_texts = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_texts.append(cell_text)
                    if row_texts:
                        text_parts.append("\t".join(row_texts))
                        text_parts.append("\n")

            extracted_text = "".join(text_parts)

            # メタデータを収集
            metadata = {
                "paragraph_count": paragraph_count,
                "table_count": table_count,
                "char_count": len(extracted_text),
            }

            # 文書のコアプロパティがあれば追加
            if hasattr(doc, "core_properties"):
                props = doc.core_properties
                if props.title:
                    metadata["title"] = props.title
                if props.author:
                    metadata["author"] = props.author
                if props.created:
                    metadata["creation_date"] = str(props.created)
                if props.modified:
                    metadata["modified_date"] = str(props.modified)

            return ExtractedText(content=extracted_text, metadata=metadata)

        except Exception as e:
            raise Exception(f"Failed to extract text from Word document: {str(e)}") from e

    def supports(self, content_type: str) -> bool:
        """指定されたコンテンツタイプをサポートしているか判定する。

        Args:
            content_type: ファイルのMIMEタイプ

        Returns:
            サポートしている場合True
        """
        base_type = content_type.split(";")[0].strip().lower()
        return base_type in self.SUPPORTED_TYPES